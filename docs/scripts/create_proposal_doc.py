"""Generate Project_Proposal_Secure_Chat_E2E.docx from proposal content."""
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE


def set_paragraph_heading(doc, text, level=1):
    """Add a heading. Level 1 = Title, 2 = Heading 1, 3 = Heading 2."""
    return doc.add_heading(text, level=level)


def add_para(doc, text, align=None):
    p = doc.add_paragraph(text)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_hyperlink_to_paragraph(paragraph, text, url):
    """Add a clickable hyperlink to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(u)
    r_pr.append(color)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_hyperlink_in_cell(cell, text, url):
    """Add a paragraph in a table cell containing a clickable hyperlink."""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()
    add_hyperlink_to_paragraph(paragraph, text, url)


def main():
    doc = Document()

    # ----- Title block (centered) -----
    add_para(doc, "Secure Chat Application with End-to-End Encryption", align="center").runs[0].bold = True
    add_para(doc, "Project Proposal", align="center").runs[0].italic = True
    doc.add_paragraph()
    add_para(doc, "Secure Chat Group", align="center")
    add_para(doc, "Zeerak Shahzad (22K-4692) · Moiz Ahmed Khan (22K-4795)", align="center")
    doc.add_paragraph()

    # ----- Contents -----
    set_paragraph_heading(doc, "Contents", level=2)
    doc.add_paragraph("1. Project Title, Group & Team", style="List Number")
    doc.add_paragraph("2. Project Description", style="List Number")
    doc.add_paragraph("3. Plan of Work (5 Weeks)", style="List Number")
    doc.add_paragraph("4. References", style="List Number")
    doc.add_paragraph()

    # ----- 1. Project Title, Group & Team -----
    set_paragraph_heading(doc, "1. Project Title, Group & Team", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Project title", "Secure Chat Application with End-to-End Encryption"),
        ("Group name", "Secure Chat Group"),
        ("Team members", "Zeerak Shahzad (22K-4692), Moiz Ahmed Khan (22K-4795)"),
    ]
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # ----- 2. Project Description -----
    set_paragraph_heading(doc, "2. Project Description", level=1)

    set_paragraph_heading(doc, "2.1 Overview", level=2)
    doc.add_paragraph(
        "We propose building a Secure Chat Application that enables private, real-time messaging with end-to-end encryption (E2E). Only the sender and the intended recipient can read messages; the server and any intermediary cannot decrypt the content. Encryption is performed on the client using keys that never leave the user's device in plain form, aligning with modern expectations for private communication."
    )
    doc.add_paragraph()

    set_paragraph_heading(doc, "2.2 Technology Stack", level=2)
    doc.add_paragraph("We use a clear separation between backend and client.")
    tech_table = doc.add_table(rows=4, cols=2)
    tech_table.style = "Table Grid"
    tech_rows = [
        ("Backend", "Node.js, Express (REST API), Socket.IO (real-time messaging), SQLite via better-sqlite3, bcryptjs (password hashing)"),
        ("Frontend", "React 19, Vite 6, TypeScript, Tailwind CSS, Motion (animations), Lucide React (icons), socket.io-client"),
        ("Cryptography", "Web Crypto API: ECDH P-256 (key agreement), AES-256-GCM (encryption), ephemeral keys per message (forward secrecy)"),
        ("Client storage", "IndexedDB (local message history); server stores only ciphertext and user/friend data"),
    ]
    for i, (layer, tech) in enumerate(tech_rows):
        tech_table.rows[i].cells[0].text = layer
        tech_table.rows[i].cells[1].text = tech
        tech_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph("The backend is API-only. The CipherChat React application serves as the client and communicates with the server over HTTP and WebSockets (HTTPS/WSS in production).")
    doc.add_paragraph()

    set_paragraph_heading(doc, "2.3 How It Will Work", level=2)
    how_items = [
        "User registration and authentication. Users sign up with a username and password. Credentials are stored securely using industry-standard hashing. Authentication is required before accessing the chat.",
        "Key generation and exchange. Each client generates a public/private key pair. Public keys are registered or exchanged via the server; private keys remain on the device. For each conversation, a shared secret is derived (e.g. using Elliptic Curve Diffie–Hellman) so that only the two participants can derive the same key.",
        "End-to-end encryption. Before sending, the client encrypts each message using the shared secret with a symmetric algorithm (AES). The ciphertext is sent to the server and forwarded to the recipient. The server never has the decryption key.",
        "Decryption on receipt. The recipient's client uses its private key and the agreed key-exchange data to derive the same shared secret and decrypt the message. Decryption happens only on the client.",
        "Secure channel and integrity. Communication with the server uses TLS (HTTPS/WSS). Message authentication (e.g. AEAD) ensures that messages are not tampered with in transit.",
    ]
    for item in how_items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    set_paragraph_heading(doc, "2.4 Functional Features", level=2)
    features = [
        "User registration. Create an account with username and password; store only a hashed password.",
        "User login / logout. Authenticate and manage session; logout clears local session and keys as appropriate.",
        "Public key registration. After login, the client generates or loads a key pair and uploads the public key to the server (the private key is never sent).",
        "Contact / user discovery. Search or list users by username and retrieve their public keys for E2E setup.",
        "Key agreement (E2E). Establish a shared secret with another user using their public key (e.g. ECDH) so that only both clients can derive the same key.",
        "Send encrypted message. Encrypt the message with the shared secret (e.g. AES), send ciphertext (and any required nonce/IV) to the server for delivery.",
        "Receive and decrypt message. Receive ciphertext from the server, derive the same shared secret on the recipient side, decrypt, and display plaintext.",
        "One-to-one chat. A dedicated channel between two users with persistent E2E encryption for that conversation.",
        "Message history (client-side). Store decrypted messages locally for the session; the server may store only ciphertext.",
        "Session and key management. Handle login session and optionally lock or clear keys on logout or after inactivity.",
    ]
    for i, f in enumerate(features, 1):
        doc.add_paragraph(f"{i}. {f}", style="List Number")
    doc.add_paragraph()

    # ----- 3. Plan of Work -----
    set_paragraph_heading(doc, "3. Plan of Work (5 Weeks)", level=1)

    set_paragraph_heading(doc, "3.1 Milestones", level=2)
    milestones_table = doc.add_table(rows=6, cols=3)
    milestones_table.style = "Table Grid"
    milestones_table.rows[0].cells[0].text = "Week"
    milestones_table.rows[0].cells[1].text = "Milestone"
    milestones_table.rows[0].cells[2].text = "Description"
    for c in range(3):
        milestones_table.rows[0].cells[c].paragraphs[0].runs[0].bold = True
    milestones_data = [
        ("1", "Requirements and design", "Finalise the functional list, architecture (client–server, key flow), and technology stack."),
        ("2", "Auth and key foundation", "Implement user registration, login/logout, and public key generation and registration."),
        ("3", "Key agreement and crypto", "Implement key agreement (e.g. ECDH) and encrypt/decrypt for one-to-one chat."),
        ("4", "Messaging pipeline", "Implement sending encrypted messages, receiving and decrypting, and the one-to-one chat channel."),
        ("5", "Integration and hardening", "Integrate all features, client-side message history and session handling, and basic testing and documentation."),
    ]
    for i, (week, milestone, desc) in enumerate(milestones_data, 1):
        milestones_table.rows[i].cells[0].text = week
        milestones_table.rows[i].cells[1].text = milestone
        milestones_table.rows[i].cells[2].text = desc
    doc.add_paragraph()

    set_paragraph_heading(doc, "3.2 Team Contributions", level=2)
    doc.add_paragraph("All functional features listed in Section 2 are implemented by the team. User interface, visual design, database design, and testing are carried out in support of these features.")
    team_table = doc.add_table(rows=3, cols=2)
    team_table.style = "Table Grid"
    team_table.rows[0].cells[0].text = "Team"
    team_table.rows[0].cells[1].text = "Zeerak Shahzad (22K-4692), Moiz Ahmed Khan (22K-4795)"
    team_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    team_table.rows[1].cells[0].text = "Responsibilities"
    team_table.rows[1].cells[1].text = (
        "User registration; login/logout; public key registration; contact and user discovery; E2E key agreement; "
        "sending encrypted messages; receiving and decrypting messages; one-to-one chat; message history and session/key management. "
        "Supporting work includes database schema for users and public keys, API endpoints for auth and key storage, "
        "server-side message relay, WebSocket API, and UI for login, registration, contacts, and chat."
    )
    team_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # ----- 4. References -----
    set_paragraph_heading(doc, "4. References", level=1)
    doc.add_paragraph("All references are provided with hyperlinks for quick access to the primary sources.")
    refs_table = doc.add_table(rows=9, cols=3)
    refs_table.style = "Table Grid"
    refs_table.rows[0].cells[0].text = "#"
    refs_table.rows[0].cells[1].text = "Reference"
    refs_table.rows[0].cells[2].text = "Link"
    for c in range(3):
        refs_table.rows[0].cells[c].paragraphs[0].runs[0].bold = True
    refs_data = [
        ("Signal Protocol Specification (Signal Foundation)", "https://signal.org/docs/", "signal.org/docs"),
        ("Messaging Layer Security (MLS), IETF", "https://datatracker.ietf.org/wg/mls/about/", "datatracker.ietf.org/wg/mls"),
        ("SubtleCrypto – Web Crypto API (MDN)", "https://developer.mozilla.org/en-US/docs/Web/API/SubtleCrypto", "MDN SubtleCrypto"),
        ("NIST SP 800-56A Rev. 3 – Key establishment", "https://csrc.nist.gov/publications/detail/sp/800-56a/rev-3/final", "NIST 800-56A"),
        ("NIST SP 800-57 Part 1 – Key management", "https://csrc.nist.gov/publications/detail/sp/800-57-part-1/rev-5/final", "NIST 800-57"),
        ("OWASP Cryptographic Storage Cheat Sheet", "https://cheatsheetseries.owasp.org/cheatsheets/Cryptographic_Storage_Cheat_Sheet.html", "OWASP Crypto Storage"),
        ("RFC 8446 – TLS 1.3 (IETF)", "https://www.rfc-editor.org/rfc/rfc8446", "RFC 8446"),
        ("OWASP Password Storage Cheat Sheet", "https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html", "OWASP Password Storage"),
    ]
    for i, (ref_title, url, link_text) in enumerate(refs_data, 1):
        refs_table.rows[i].cells[0].text = str(i)
        refs_table.rows[i].cells[1].text = ref_title
        refs_table.rows[i].cells[2].paragraphs[0].clear()
        add_hyperlink_to_paragraph(refs_table.rows[i].cells[2].paragraphs[0], link_text, url)

    out_path = os.path.join(os.path.dirname(__file__), "Project_Proposal_Secure_Chat_E2E.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
